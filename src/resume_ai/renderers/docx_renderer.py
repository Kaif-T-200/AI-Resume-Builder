from pathlib import Path

from resume_ai.models import Resume

try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None


class DocxRenderer:
    def __init__(self):
        if Document is None:
            raise ImportError("python-docx is required for DOCX rendering; install with `pip install python-docx`")

    def render(self, resume: Resume, *, output_path: str) -> None:
        document = Document()
        document.add_heading(resume.contact.full_name, level=0)

        if resume.summary:
            document.add_heading("Summary", level=1)
            document.add_paragraph(resume.summary)

        if resume.experience:
            document.add_heading("Experience", level=1)
            for exp in resume.experience:
                p = document.add_paragraph()
                p.add_run(f"{exp.title} at {exp.company or 'N/A'}").bold = True
                if exp.bullets:
                    for bullet in exp.bullets:
                        document.add_paragraph(bullet, style="List Bullet")

        if resume.education:
            document.add_heading("Education", level=1)
            for edu in resume.education:
                document.add_paragraph(f"{edu.degree or ''} - {edu.institution}")

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_file))
